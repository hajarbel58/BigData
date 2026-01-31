import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd
from docx import Document
from docx.shared import Inches  # or from docx import Inches
import os

# Ecrivons les données de l'output d'analyse de sentiment
sentiments = {'Negative': 56378, 'Positive': 64516, 'Neutral': 28728}

# Convertissons les données en DataFrame pour une meilleure manipulation
data = pd.DataFrame({'Sentiment': ['Negative', 'Positive', 'Neutral'], 'Count': [56378, 64516, 28728]})

# Generons maintenant et sauvegardons les graphiques

# Diagramme en barres
def generate_bar_chart():
    plt.figure(figsize=(8, 5))
    sns.barplot(x=data['Sentiment'], y=data['Count'], palette=['red', 'green','gray'])
    plt.title("Analyse de Sentiment - Nombre de Textes")
    plt.ylabel("Nombre")
    plt.xlabel("Sentiments")
    plt.savefig("sentiment_barplot.png")
    plt.close()

# Diagramme circulaire (Pie Chart)
def generate_pie_chart():
    plt.figure(figsize=(7, 7))
    plt.pie(data['Count'], labels=data['Sentiment'], autopct='%1.1f%%', colors=['red', 'green','gray'], startangle=90)
    plt.title("Répartition des Sentiments")
    plt.savefig("sentiment_piechart.png")
    plt.close()

# Diagramme des pourcentages
def generate_percentage_chart():
    total = data['Count'].sum()
    data['Percentage'] = (data['Count'] / total) * 100

    plt.figure(figsize=(8, 5))
    sns.barplot(data=data, x='Percentage', y='Sentiment', palette=['red', 'green','gray'])
    plt.title("Pourcentage des Sentiments")
    plt.xlabel("Pourcentage")
    plt.ylabel("Sentiments")
    plt.xlim(0, 100)
    plt.savefig("sentiment_percentage.png")
    plt.close()

# Boxplot (Boîte à moustaches)
def generate_boxplot():
    plt.figure(figsize=(8, 5))
    sns.boxplot(data=data, x='Sentiment', y='Count', palette=['red', 'green','gray'])
    plt.title("Analyse par Boîte à Moustaches des Sentiments")
    plt.ylabel("Nombre de Textes")
    plt.xlabel("Sentiments")
    plt.savefig("sentiment_boxplot.png")
    plt.close()

# Créeons un fichier Word et insérons les graphiques
def create_word_report():
    doc = Document()
    doc.add_heading("Rapport d'Analyse de Sentiment", level=1)

    # Ajoutons les graphiques au fichier Word
    charts = [
        ("sentiment_barplot.png", "Diagramme en Barres"),
        ("sentiment_piechart.png", "Diagramme Circulaire"),
        ("sentiment_percentage.png", "Pourcentage des Sentiments"),
        ("sentiment_boxplot.png", "Analyse par Boîte à Moustaches")
    ]

    for file_name, title in charts:
        if os.path.exists(file_name):
            doc.add_heading(title, level=2)
            doc.add_picture(file_name, width=Inches(5))
        else:
            print(f"Erreur : Le fichier {file_name} n'existe pas.")

    # Enregistrons le document Word
    doc.save("sentiment_analysis_report.docx")
    print("Rapport Word généré : sentiment_analysis_report.docx")

# Executons les  fonctions
generate_bar_chart()
generate_pie_chart()
generate_percentage_chart()
generate_boxplot()

# Créeons le rapport Word avec les graphiques deja enregistrés
create_word_report()

